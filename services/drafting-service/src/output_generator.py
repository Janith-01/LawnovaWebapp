from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from config import FONTS_DIR, OUTPUT_DIR


TITLE_TEXT = {
    ("AFFIDAVIT", "en"): "AFFIDAVIT",
    ("AFFIDAVIT", "si"): "දිවුරුම් ප්‍රකාශය",
    ("CONTRACT", "en"): "AGREEMENT",
    ("CONTRACT", "si"): "ගිවිසුම",
    ("PETITION", "en"): "PETITION",
    ("PETITION", "si"): "පෙත්සම",
}

EN_SIGNATURE_PREFIXES = (
    "Signature",
    "Before me",
    "Justice of the Peace",
    "Name:",
    "Date:",
    "Address:",
    "Designation:",
    "Witnesses:",
    "Name of Signatory:",
    "Address for Service:",
)

SI_SIGNATURE_PREFIXES = (
    "ප්‍රකාශකයාගේ අත්සන",
    "මා ඉදිරියේ",
    "සාම විනිසුරුවරයා",
    "දිවුරුම් කොමසාරිස්",
    "නම:",
    "දිනය:",
    "ලිපිනය:",
    "තනතුර:",
    "සාක්ෂිකරුවන්:",
    "අත්සන්කරුගේ නම:",
    "සේවා සඳහා ලිපිනය:",
)


def _ensure_output_dir() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def _timestamped_path(doc_type: str, language: str, suffix: str) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{doc_type}_{language}_{timestamp}.{suffix}"
    return _ensure_output_dir() / filename


def _normalize_content(content: str) -> str:
    text = (content or "").replace("\r\n", "\n").replace("\r", "\n")
    return text.strip()


def _docx_font_name(language: str) -> str:
    if language == "en":
        return "Times New Roman"

    noto_font = FONTS_DIR / "NotoSansSinhala-Regular.ttf"
    if noto_font.exists():
        return "Noto Sans Sinhala"
    return "Nirmala UI"


def _set_run_font(run, font_name: str, size: int, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _count_indent(raw_line: str) -> int:
    return len(raw_line) - len(raw_line.lstrip(" "))


def _is_title_line(line: str, doc_type: str, language: str) -> bool:
    expected_title = TITLE_TEXT.get((doc_type, language), "")
    return line.strip() == expected_title


def _is_signature_line(line: str, language: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False

    if stripped.startswith(("................................", "________________")):
        return True

    prefixes = EN_SIGNATURE_PREFIXES if language == "en" else SI_SIGNATURE_PREFIXES
    return any(stripped.startswith(prefix) for prefix in prefixes)


def _is_heading_line(line: str, doc_type: str, language: str) -> bool:
    stripped = line.strip()
    if not stripped or _is_title_line(stripped, doc_type, language):
        return False

    if language == "en":
        if stripped.isupper() and len(stripped) <= 120:
            return True
        if stripped.startswith(("IN THE ", "TO: ", "NOW IT IS HEREBY AGREED", "IN WITNESS WHEREOF")):
            return True
        if stripped in {"Between", "And", "Petitioner", "Respondent"}:
            return True
    else:
        if stripped.startswith(("ශ්‍රී ලංකා", "ගරු අධිකරණය වෙත", "එබැවින්", "අතර", "සහ")):
            return True
        if stripped in {"පාර්ශවයන්", "පූර්වකථනය", "පෙත්සම", "පෙත්සම්කරු", "විත්තිකරු"}:
            return True

    return False


def _create_docx_paragraph(document: Document, line: str, style_type: str, font_name: str) -> None:
    paragraph = document.add_paragraph()
    indent_points = _count_indent(line) * 2
    text = line.strip()
    run = paragraph.add_run(text)

    if style_type == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.0
        _set_run_font(run, font_name, 14, bold=True)
        return

    if style_type == "heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.isupper() or text.startswith(("IN THE ", "ශ්‍රී ලංකා")) else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Pt(indent_points)
        _set_run_font(run, font_name, 12, bold=True)
        return

    paragraph.paragraph_format.left_indent = Pt(indent_points)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)

    if style_type == "signature":
        paragraph.paragraph_format.space_before = Pt(6)
        _set_run_font(run, font_name, 11, bold=False)
    else:
        _set_run_font(run, font_name, 11, bold=False)


def _populate_docx(document: Document, content: str, doc_type: str, language: str) -> None:
    font_name = _docx_font_name(language)
    lines = _normalize_content(content).splitlines()
    previous_was_blank = True

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for raw_line in lines:
        if not raw_line.strip():
            document.add_paragraph()
            previous_was_blank = True
            continue

        stripped = raw_line.strip()
        if _is_title_line(stripped, doc_type, language):
            style_type = "title"
        elif _is_heading_line(stripped, doc_type, language):
            style_type = "heading"
        elif _is_signature_line(stripped, language):
            style_type = "signature"
        else:
            style_type = "body"

        if style_type == "signature" and not previous_was_blank:
            document.add_paragraph()

        _create_docx_paragraph(document, raw_line, style_type, font_name)
        previous_was_blank = False


def _register_pdf_fonts(language: str) -> dict:
    if language == "en":
        return {"regular": "Times-Roman", "bold": "Times-Bold"}

    regular_name = "LawnovaSinhalaRegular"
    bold_name = "LawnovaSinhalaBold"

    preferred_regular = FONTS_DIR / "NotoSansSinhala-Regular.ttf"
    preferred_bold = FONTS_DIR / "NotoSansSinhala-Bold.ttf"

    if preferred_regular.exists():
        try:
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(preferred_regular)))
            if preferred_bold.exists():
                if bold_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(bold_name, str(preferred_bold)))
            else:
                bold_name = regular_name
            return {"regular": regular_name, "bold": bold_name}
        except Exception:
            pass

    for candidate in sorted(FONTS_DIR.glob("*.ttf")) + sorted(FONTS_DIR.glob("*.otf")):
        try:
            fallback_name = f"LawnovaFallback_{candidate.stem}"
            if fallback_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(fallback_name, str(candidate)))
            return {"regular": fallback_name, "bold": fallback_name}
        except Exception:
            continue

    return {"regular": "Helvetica", "bold": "Helvetica-Bold"}


def _get_pdf_style(base_style: ParagraphStyle, name: str, font_name: str, font_size: int, alignment: int, left_indent: int = 0) -> ParagraphStyle:
    return ParagraphStyle(
        name,
        parent=base_style,
        fontName=font_name,
        fontSize=font_size,
        leading=font_size + 3,
        alignment=alignment,
        leftIndent=left_indent,
        spaceAfter=6,
        spaceBefore=0,
    )


def _build_pdf_story(content: str, doc_type: str, language: str, fonts: dict) -> list:
    styles = getSampleStyleSheet()
    title_style = _get_pdf_style(styles["Normal"], "DraftTitle", fonts["bold"], 14, TA_CENTER)
    heading_style = _get_pdf_style(styles["Normal"], "DraftHeading", fonts["bold"], 12, TA_LEFT)
    centered_heading_style = _get_pdf_style(styles["Normal"], "DraftHeadingCentered", fonts["bold"], 12, TA_CENTER)
    signature_style = _get_pdf_style(styles["Normal"], "DraftSignature", fonts["regular"], 11, TA_LEFT)
    body_style = _get_pdf_style(styles["Normal"], "DraftBody", fonts["regular"], 11, TA_LEFT)
    story = []
    previous_was_blank = True

    for raw_line in _normalize_content(content).splitlines():
        if not raw_line.strip():
            story.append(Spacer(1, 6))
            previous_was_blank = True
            continue

        stripped = raw_line.strip()
        indent_points = _count_indent(raw_line) * 2
        safe_text = escape(stripped)

        if _is_title_line(stripped, doc_type, language):
            style = title_style
        elif _is_heading_line(stripped, doc_type, language):
            style = centered_heading_style if stripped.isupper() or stripped.startswith(("IN THE ", "ශ්‍රී ලංකා")) else heading_style
        elif _is_signature_line(stripped, language):
            style = _get_pdf_style(styles["Normal"], f"DraftSignature{indent_points}", fonts["regular"], 11, TA_LEFT, indent_points)
            if not previous_was_blank:
                story.append(Spacer(1, 8))
        else:
            style = _get_pdf_style(styles["Normal"], f"DraftBody{indent_points}", fonts["regular"], 11, TA_LEFT, indent_points)

        story.append(Paragraph(safe_text, style))
        story.append(Spacer(1, 4))
        previous_was_blank = False

    return story


def _build_pdf_document(output_path: Path, story: list) -> None:
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )
    pdf.build(story)


def _build_canvas_fallback(output_path: Path, content: str, font_name: str) -> None:
    canvas = Canvas(str(output_path), pagesize=A4)
    page_width, page_height = A4
    y_position = page_height - 54
    max_width = page_width - 108
    text_object = canvas.beginText(54, y_position)
    text_object.setFont(font_name, 11)
    text_object.setLeading(14)

    for line in _normalize_content(content).splitlines():
        if y_position <= 54:
            canvas.drawText(text_object)
            canvas.showPage()
            y_position = page_height - 54
            text_object = canvas.beginText(54, y_position)
            text_object.setFont(font_name, 11)
            text_object.setLeading(14)

        safe_line = line.rstrip()
        if len(safe_line) > 140:
            while safe_line:
                segment = safe_line[:140]
                safe_line = safe_line[140:]
                text_object.textLine(segment)
                y_position -= 14
        else:
            text_object.textLine(safe_line)
            y_position -= 14

    canvas.drawText(text_object)
    canvas.save()


def save_as_docx(content: str, doc_type: str, language: str) -> str:
    output_path = _timestamped_path(doc_type, language, "docx")

    try:
        document = Document()
        _populate_docx(document, content, doc_type, language)
        document.save(output_path)
    except Exception:
        fallback_document = Document()
        fallback_font = _docx_font_name(language)
        for block in _normalize_content(content).split("\n\n"):
            paragraph = fallback_document.add_paragraph()
            run = paragraph.add_run(block.strip())
            _set_run_font(run, fallback_font, 11, bold=False)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)
        fallback_document.save(output_path)

    return str(output_path)


def save_as_pdf(content: str, doc_type: str, language: str) -> str:
    output_path = _timestamped_path(doc_type, language, "pdf")
    fonts = _register_pdf_fonts(language)

    try:
        story = _build_pdf_story(content, doc_type, language, fonts)
        _build_pdf_document(output_path, story)
    except Exception:
        fallback_font = fonts["regular"] if fonts.get("regular") else "Helvetica"
        try:
            _build_canvas_fallback(output_path, content, fallback_font)
        except Exception:
            _build_canvas_fallback(output_path, content, "Helvetica")

    return str(output_path)
